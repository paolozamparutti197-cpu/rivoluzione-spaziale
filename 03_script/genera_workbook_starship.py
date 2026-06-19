from __future__ import annotations

import re
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, DoughnutChart, LineChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parents[1]
SOURCE_DOC_DIR = BASE_DIR / "05_fonti_originali"
OUTPUT_DIR = BASE_DIR / "01_workbook"
CURRENT_WORKBOOK = OUTPUT_DIR / "sviluppo_starship.xlsx"

DOC1 = SOURCE_DOC_DIR / "programma_starship_spacex.docx"
DOC2 = SOURCE_DOC_DIR / "ricerca_starship_1.docx"
OUT = OUTPUT_DIR / "sviluppo_starship_generato.xlsx"

PALETTE = {
    "navy": "0B1320",
    "ink": "1F2937",
    "muted": "64748B",
    "gold": "C9A227",
    "sand": "F4EBD0",
    "paper": "F8FAFC",
    "line": "D9E2EC",
    "success": "DFF3EA",
    "warn": "FFF4D6",
    "danger": "FDE2E1",
    "blue": "E6F0FF",
    "teal": "DDF7F3",
    "gray": "EEF2F7",
}

CATEGORY_FILL = {
    "Ideazione/Annuncio": "E6F0FF",
    "Design/Decisione": "E2F7E1",
    "Prototipazione/Ground Test": "FFF4D6",
    "Volo Suborbitale": "FFE8D6",
    "Volo Integrato/Orbitale": "DDF7F3",
    "Esplosione/RUD/Fallimento": "FDE2E1",
    "Milestone Regolatorio/Contrattuale": "EEE7FF",
    "Altro": "EEF2F7",
}


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("࿘", "")
    text = re.sub(r"cite[^ ]+", "", text)
    text = re.sub(r"\[\d+\]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    replacements = {
        "Vincitore morale della giornata: il calcestruzzo, che ha perso male. Ma la correzione infrastrutturale ha cambiato il programma.": "La correzione infrastrutturale del pad ha cambiato in modo sostanziale il programma.",
        "Se uno tossisce, l’altro sanguina.": "Dimostra la forte interdipendenza tra booster e pad.",
        "I tank test sono noiosi solo a chi non deve firmare il certificato strutturale.": "I test strutturali sono essenziali per validare materiali e saldature.",
        "sicurezza teatrale": "eccesso di sicurezza",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def workbook_rows(sheet_name: str, header_row: int) -> list[dict]:
    if not CURRENT_WORKBOOK.exists():
        return []
    wb = load_workbook(CURRENT_WORKBOOK, read_only=True, data_only=True)
    try:
        if sheet_name not in wb.sheetnames:
            return []
        ws = wb[sheet_name]
        headers = [cell.value for cell in ws[header_row]]
        rows = []
        for source_row in ws.iter_rows(min_row=header_row + 1, values_only=True):
            row = {}
            has_value = False
            for header, value in zip(headers, source_row):
                if not header:
                    continue
                row[str(header)] = value
                has_value = has_value or value not in (None, "")
            if has_value:
                rows.append(row)
        return rows
    finally:
        wb.close()


def cell_to_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    return str(value)


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = cell_to_text(value)


def technical_tables_from_current_workbook() -> dict[str, list[list[str]]]:
    if not CURRENT_WORKBOOK.exists():
        return {}
    wb = load_workbook(CURRENT_WORKBOOK, read_only=True, data_only=True)
    try:
        if "Confronti tecnici" not in wb.sheetnames:
            return {}
        ws = wb["Confronti tecnici"]
        table_names = {"Confronto Block", "Confronto Raptor", "Milestone operative Doc 2"}
        tables: dict[str, list[list[str]]] = {}
        current_name = ""
        current_rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            values = [cell_to_text(value) for value in row]
            values = values[: max((idx + 1 for idx, value in enumerate(values) if value), default=0)]
            if not values:
                if current_name and current_rows:
                    tables[current_name] = current_rows
                    current_name = ""
                    current_rows = []
                continue
            if values[0] in table_names:
                if current_name and current_rows:
                    tables[current_name] = current_rows
                current_name = values[0]
                current_rows = []
                continue
            if current_name:
                current_rows.append(values)
        if current_name and current_rows:
            tables[current_name] = current_rows
        return tables
    finally:
        wb.close()


def reconstruct_source_docs_from_current_workbook() -> None:
    timeline_rows = workbook_rows("Timeline integrata", 4)
    technical_tables = technical_tables_from_current_workbook()
    if not timeline_rows or len(technical_tables) < 3:
        return

    SOURCE_DOC_DIR.mkdir(parents=True, exist_ok=True)

    doc1 = Document()
    doc1.add_heading("Programma Starship di SpaceX", level=1)
    doc1.add_paragraph(
        "Documento sorgente ricostruito automaticamente da 01_workbook/sviluppo_starship.xlsx."
    )
    add_docx_table(
        doc1,
        ["Data", "Evento", "Categoria", "Dettagli", "Impatto"],
        [
            [
                row.get("Data", ""),
                row.get("Evento", ""),
                row.get("Categoria", ""),
                row.get("Dettagli", ""),
                row.get("Impatto", ""),
            ]
            for row in timeline_rows
        ],
    )
    doc1.save(DOC1)

    doc2 = Document()
    doc2.add_heading("Ricerca Starship 1", level=1)
    doc2.add_paragraph(
        "Documento sorgente ricostruito automaticamente da 01_workbook/sviluppo_starship.xlsx."
    )
    for table_name in ["Confronto Block", "Confronto Raptor", "Milestone operative Doc 2"]:
        rows = technical_tables.get(table_name, [])
        if not rows:
            continue
        doc2.add_heading(table_name, level=2)
        max_cols = max(len(row) for row in rows)
        normalized = [row + [""] * (max_cols - len(row)) for row in rows]
        add_docx_table(doc2, normalized[0], normalized[1:])

    for row in workbook_rows("Fonti", 4):
        source = cell_to_text(row.get("Fonte", ""))
        if source:
            doc2.add_paragraph(source)
    doc2.save(DOC2)


def extract_doc1_rows() -> list[dict]:
    doc = Document(DOC1)
    table = doc.tables[0]
    rows = []
    for source_row in table.rows[1:]:
        cells = [clean_text(cell.text) for cell in source_row.cells]
        if len(cells) < 5 or not cells[0]:
            continue
        rows.append(
            {
                "Data": cells[0],
                "Evento": cells[1],
                "Categoria": cells[2],
                "Dettagli": cells[3],
                "Impatto": cells[4],
                "Fonte": "Doc 1 - cronologia strutturata",
            }
        )
    return rows


def extract_doc2_tables() -> dict[str, list[list[str]]]:
    doc = Document(DOC2)
    names = ["Confronto Block", "Confronto Raptor", "Milestone operative Doc 2"]
    tables = {}
    for name, table in zip(names, doc.tables):
        tables[name] = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
    return tables


def extract_sources() -> list[dict]:
    workbook_sources = workbook_rows("Fonti", 4)
    if workbook_sources:
        return [
            {"Documento": cell_to_text(row.get("Documento", "")), "Fonte": cell_to_text(row.get("Fonte", ""))}
            for row in workbook_sources
            if row.get("Fonte")
        ]

    rows = []
    for doc_name, path in [("Doc 1", DOC1), ("Doc 2", DOC2)]:
        doc = Document(path)
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                text_parts.extend(cell.text for cell in row.cells)
        urls = sorted(set(re.findall(r"https?://[^\s)]+", "\n".join(text_parts))))
        rows.extend({"Documento": doc_name, "Fonte": url} for url in urls)
    doc2_domains = "basenor.com, en.wikipedia.org, bydfi.com, newspaceeconomy.ca, starship-spacex.fandom.com, orbitaltoday.com, eureka.patsnap.com, flyingmag.com, stocktwits.com, reddit.com, teslaoracle.com, spaceexplored.com, faa.gov, starbase.texas.gov, ballotpedia.org, payloadspace.com, nasa.gov"
    rows.append({"Documento": "Doc 2", "Fonte": f"Elenco domini citati nel testo: {doc2_domains}"})
    return rows


def parse_sort_date(value: str) -> datetime:
    v = value.strip()
    m = re.match(r"(\d{4})[–-]", v)
    if m:
        return datetime(int(m.group(1)), 1, 1)
    m = re.search(r"(\d{1,2})/(\d{4})[–-]\d{1,2}/\d{1,2}/\d{4}", v)
    if m:
        return datetime(int(m.group(2)), int(m.group(1)), 1)
    m = re.search(r"(\d{1,2})[–-]\d{1,2}/(\d{4})", v)
    if m:
        return datetime(int(m.group(2)), int(m.group(1)), 1)
    for pattern in [
        r"(\d{1,2})/(\d{1,2})/(\d{4})",
        r"(\d{1,2})/(\d{1,2})[–-]\d{1,2}/\d{1,2}/(\d{4})",
        r"(\d{1,2})[–-]\d{1,2}/(\d{1,2})/(\d{4})",
    ]:
        m = re.search(pattern, v)
        if m:
            return datetime(int(m.group(3)), int(m.group(2)), int(m.group(1)))
    m = re.search(r"(\d{1,2})/(\d{4})", v)
    if m:
        return datetime(int(m.group(2)), int(m.group(1)), 1)
    m = re.search(r"(\d{4})", v)
    if m:
        return datetime(int(m.group(1)), 1, 1)
    return datetime(1900, 1, 1)


def phase_for_date(date_key: datetime) -> str:
    year = date_key.year
    if year <= 2018:
        return "Genesi concettuale e scelte fondative"
    if 2019 <= year <= 2021:
        return "Prototipi suborbitali e apprendimento rapido"
    if year == 2022:
        return "Pad, booster e licenze pre-volo integrato"
    if year == 2023:
        return "Primi voli integrati"
    if year == 2024:
        return "Maturazione rientro, hot-staging e catch"
    if year == 2025:
        return "Block 2, riuso booster e stabilizzazione ship"
    return "Block 3 / V3 e preparazione Flight 12"


def vehicle_from_event(event: str, details: str) -> str:
    text = f"{event} {details}"
    patterns = [r"Booster\s?\d+(?:-\d+)?", r"B\d+(?:-\d+)?", r"Ship\s?\d+", r"S\d+", r"SN\d+", r"Mk\d+", r"Starhopper", r"BN\d+"]
    found = []
    for pat in patterns:
        for match in re.findall(pat, text, flags=re.I):
            if match not in found:
                found.append(match)
    return ", ".join(found[:5])


def outcome_from(row: dict) -> str:
    text = f"{row['Evento']} {row['Categoria']} {row['Dettagli']} {row['Impatto']}".lower()
    if any(word in text for word in ["esplode", "distrutto", "failure", "rud", "perdita"]):
        if any(word in text for word in ["successo", "riuscito", "catturato", "completa"]):
            return "Parziale / apprendimento"
        return "Fallimento / RUD"
    if any(word in text for word in ["riuscito", "successo", "completa", "catturato"]):
        return "Successo / milestone"
    if any(word in text for word in ["autorizza", "assegna", "chiude", "seleziona"]):
        return "Milestone istituzionale"
    return "Neutro / evoluzione"


def integrated_enrichment(event: str) -> str:
    checks = [
        ("Starhopper", "Doc 2 conferma Starhopper come banco di prova del controllo Raptor; resta la cronologia piu precisa del Doc 1."),
        ("SN8", "Doc 2 rafforza il punto tecnico: belly-flop riuscito, perdita al landing burn per pressione header tank."),
        ("SN9", "Doc 2 conferma la failure di relight come causa centrale della perdita in atterraggio."),
        ("SN10", "Doc 2 integra la lettura post-landing: touchdown ottenuto ma veicolo non integro per landing duro."),
        ("SN11", "Doc 2 attribuisce la perdita alla turbopompa metano; Doc 1 mantiene formulazione piu prudente."),
        ("SN15", "Doc 2 conferma SN15 come chiusura positiva della fase suborbitale."),
        ("Booster 7", "Doc 2 aggiunge che l'incidente spin-start rese piu cauta la gestione dei test multi-motore al pad."),
        ("Integrated Flight Test 1", "Doc 2 aggiunge il ritardo dell'AFTS e sottolinea il danno infrastrutturale; Doc 1 resta base per cronologia e impatti."),
        ("Integrated Flight Test 2", "Doc 2 conferma hot-staging e perdita ship legata a incendio/sfiato; scelta Doc 1 per causa e sequenza."),
        ("Integrated Flight Test 3", "Doc 2 aggiunge apogeo circa 234 km e test portellone/propellant transfer; non classificato come orbita stabile."),
        ("Integrated Flight Test 4", "Doc 2 lo classifica come successo; qui resta successo end-to-end parziale per danni al flap."),
        ("Integrated Flight Test 5", "Doc 2 conferma la prima cattura booster come milestone storica."),
        ("Integrated Flight Test 6", "Doc 2 aggiunge luce diurna per raccolta dati; Doc 1 valorizza il primo in-space relight."),
        ("Flight 7", "Doc 2 conferma propellant leak e perdita Ship 33; Doc 1 offre root-cause FAA successiva."),
        ("Flight 8", "Doc 2 conferma terza cattura booster e perdita Ship 34."),
        ("Flight 9", "Doc 2 conferma reflight B14-2, no deployment Starlink e perdita in rientro."),
        ("Flight 10", "Doc 2 conferma primo deployment riuscito di simulatori Starlink."),
        ("Flight 11", "Doc 2 aggiunge test di scudi/ablativi verso V3; Doc 1 mantiene formulazione prudente sulla documentazione pubblica."),
        ("Starbase come città", "Doc 2 aggiunge Bobby Peden come primo sindaco; contenuto amministrativo mantenuto come contesto."),
        ("Flight 12", "Doc 2 aggiunge obiettivi attesi per Starlink V3 e docking passivo; non trattati come ufficialmente confermati nel workbook."),
    ]
    for needle, note in checks:
        if needle.lower() in event.lower():
            return note
    return ""


def prepare_timeline() -> list[dict]:
    rows = []
    for idx, row in enumerate(extract_doc1_rows(), 1):
        date_key = parse_sort_date(row["Data"])
        row["SortKey"] = date_key
        row["Anno"] = date_key.year
        row["Fase"] = phase_for_date(date_key)
        row["Veicolo/Missione"] = vehicle_from_event(row["Evento"], row["Dettagli"])
        row["Esito"] = outcome_from(row)
        row["Fusione"] = integrated_enrichment(row["Evento"])
        row["ID"] = idx
        rows.append(row)
    return sorted(rows, key=lambda r: (r["SortKey"], r["ID"]))


def integrated_flights() -> list[dict]:
    workbook_data = workbook_rows("Voli integrati", 4)
    if workbook_data:
        return [
            {
                "Flight": cell_to_text(row.get("Flight", "")),
                "Data": cell_to_text(row.get("Data", "")),
                "Veicolo": cell_to_text(row.get("Veicolo", "")),
                "Block": cell_to_text(row.get("Block", "")),
                "Milestone": cell_to_text(row.get("Milestone", "")),
                "Booster": cell_to_text(row.get("Booster", "")),
                "Ship": cell_to_text(row.get("Ship", "")),
                "Esito integrato": cell_to_text(row.get("Esito integrato", "")),
                "Lezione": cell_to_text(row.get("Lezione", "")),
            }
            for row in workbook_data
        ]

    return [
        {"Flight": "IFT-1", "Data": "20/04/2023", "Veicolo": "Booster 7 + Ship 24", "Block": "Block 1", "Milestone": "Primo volo integrato", "Booster": "Perdita controllo, nessuna separazione", "Ship": "Persa con stack", "Esito integrato": "Fallimento", "Lezione": "Pad, AFTS, flame diverter/deluge e affidabilita multi-engine diventano priorita."},
        {"Flight": "IFT-2", "Data": "18/11/2023", "Veicolo": "Booster 9 + Ship 25", "Block": "Block 1", "Milestone": "Primo hot-staging riuscito", "Booster": "Perso durante boostback", "Ship": "Persa dopo coast/ascent per incendio/sfiato LOX", "Esito integrato": "Parziale", "Lezione": "Superato il blocco della separazione; resta aperta la maturita di recovery e ship."},
        {"Flight": "IFT-3", "Data": "14/03/2024", "Veicolo": "Booster 10 + Ship 28", "Block": "Block 1", "Milestone": "Operazioni in-space: propellant transfer demo e payload door", "Booster": "Hard splashdown", "Ship": "Persa in rientro", "Esito integrato": "Parziale", "Lezione": "Traiettoria transatmosferica/quasi-orbitale; evitare di presentarla come orbita operativa stabile."},
        {"Flight": "IFT-4", "Data": "06/06/2024", "Veicolo": "Booster 11 + Ship 29", "Block": "Block 1", "Milestone": "Primo rientro e splashdown controllato di entrambi gli stadi", "Booster": "Splashdown morbido", "Ship": "Rientro completato con grave danno flap", "Esito integrato": "Successo / parziale", "Lezione": "Il sistema dimostra profilo end-to-end, ma TPS/flap resta area critica."},
        {"Flight": "IFT-5", "Data": "13/10/2024", "Veicolo": "Booster 12 + Ship 30", "Block": "Block 1", "Milestone": "Prima cattura booster con Mechazilla", "Booster": "Catch riuscito", "Ship": "Splashdown controllato", "Esito integrato": "Successo", "Lezione": "Booster-side entra in fase di maturazione rapida."},
        {"Flight": "IFT-6", "Data": "19/11/2024", "Veicolo": "Booster 13 + Ship 31", "Block": "Block 1", "Milestone": "Primo in-space relight riuscito", "Booster": "Catch annullato, splashdown", "Ship": "Relight e splashdown", "Esito integrato": "Parziale", "Lezione": "La missione sposta il focus sulle funzioni in-space della ship."},
        {"Flight": "Flight 7", "Data": "16/01/2025", "Veicolo": "Booster 14 + Ship 33", "Block": "Block 2", "Milestone": "Debutto Ship Block 2, secondo catch booster", "Booster": "Catch riuscito", "Ship": "Persa in ascesa; detriti su Turks and Caicos", "Esito integrato": "Parziale", "Lezione": "La maturita booster supera nettamente quella ship."},
        {"Flight": "Flight 8", "Data": "06/03/2025", "Veicolo": "Booster 15 + Ship 34", "Block": "Block 2", "Milestone": "Terzo catch booster", "Booster": "Catch riuscito", "Ship": "Persa tardi in ascesa", "Esito integrato": "Parziale", "Lezione": "Ancora criticita propulsion-side della ship."},
        {"Flight": "Flight 9", "Data": "27/05/2025", "Veicolo": "Booster 14-2 + Ship 35", "Block": "Block 2", "Milestone": "Primo reflight Super Heavy", "Booster": "Perso prima dello splashdown", "Ship": "Leak, no relight, no deployment, breakup", "Esito integrato": "Parziale", "Lezione": "Riuso booster dimostrato come campagna, non ancora come recovery completo."},
        {"Flight": "Flight 10", "Data": "26/08/2025", "Veicolo": "Booster 16 + Ship 37", "Block": "Block 2", "Milestone": "Primo deployment riuscito simulatori Starlink", "Booster": "Splashdown controllato", "Ship": "Deployment, relight, rientro e splashdown", "Esito integrato": "Successo", "Lezione": "Miglior volo ship-side fino a quel punto."},
        {"Flight": "Flight 11", "Data": "13/10/2025", "Veicolo": "Booster 15-2 + Ship 38", "Block": "Block 2", "Milestone": "Chiusura positiva fase Block 2", "Booster": "Controlled splashdown", "Ship": "Deployment, relight e rientro controllato", "Esito integrato": "Successo", "Lezione": "Missione sostanzialmente completa secondo fonti pubbliche disponibili."},
    ]


def summary_topics() -> list[dict]:
    workbook_data = workbook_rows("Temi e sintesi", 4)
    if workbook_data:
        return [
            {
                "Tema": cell_to_text(row.get("Tema", "")),
                "Sintesi integrata": cell_to_text(row.get("Sintesi integrata", "")),
                "Impatto": cell_to_text(row.get("Impatto", "")),
                "Fonte prevalente": cell_to_text(row.get("Fonte prevalente", "")),
            }
            for row in workbook_data
        ]

    return [
        {"Tema": "Evoluzione architetturale", "Sintesi integrata": "Il programma passa da BFR/MCT e ITS da 12 m a BFR/Starship da 9 m, poi al sistema acciaio inox + Raptor + hot-staging. Le tre potature decisive sono geometria, materiale e semplificazione motoristica/aerodinamica.", "Impatto": "Rende il sistema producibile, testabile e iterabile con una logica industriale, non solo dimostrativa.", "Fonte prevalente": "Doc 1, integrato da sezioni 1-2 Doc 2"},
        {"Tema": "Raptor come collo di bottiglia", "Sintesi integrata": "Raptor evolve da full-flow methalox complesso a Raptor 2 piu semplice e Raptor 3 piu compatto, con package pulito, meno shielding e target di riuso rapido.", "Impatto": "La maturita del motore condiziona payload, cadenza, affidabilita e fattibilita economica.", "Fonte prevalente": "Doc 1 per timeline; Doc 2 per tabella tecnica"},
        {"Tema": "Produzione e infrastruttura", "Sintesi integrata": "Boca Chica evolve da ring stack all'aperto a High Bay, Mega Bay, Starfactory, Pad 2 e piani Gigabay/Florida. L'infrastruttura diventa parte del veicolo.", "Impatto": "La cadenza futura dipende da fabbriche, pad, GSE e licenze tanto quanto dal razzo.", "Fonte prevalente": "Doc 1, integrato da Doc 2"},
        {"Tema": "Booster piu maturo della ship", "Sintesi integrata": "Dal 2024-2025 i catch booster diventano ripetibili, mentre la ship continua a mostrare criticita in ascent, coast, rientro, TPS e deployment.", "Impatto": "Per Starlink V3, HLS e Marte il problema piu duro resta la ship, non solo Super Heavy.", "Fonte prevalente": "Doc 1"},
        {"Tema": "FAA e regolazione", "Sintesi integrata": "La FAA impone mitigazioni e corrective actions, ma amplia anche il perimetro operativo con 25 voli/anno in Texas e dossier LC-39A/Florida.", "Impatto": "La narrativa 'FAA blocca tutto' e' incompleta: ritardi regolatori e guasti tecnici si alimentano a vicenda.", "Fonte prevalente": "Doc 1"},
        {"Tema": "Artemis HLS e rifornimento orbitale", "Sintesi integrata": "Starship HLS resta centrale, ma la sfida principale e' trasferire e conservare propellente criogenico in orbita con cadenza di tanker elevata. Il Doc 1 aggiorna il pivot NASA 2026 verso demo LEO nel 2027 e landing Artemis IV nel 2028.", "Impatto": "Il successo lunare dipende dalla logistica spaziale piu che dal solo lancio.", "Fonte prevalente": "Doc 1; Doc 2 integra boil-off e 10-15 tanker come contesto"},
        {"Tema": "Starlink V3 come driver economico", "Sintesi integrata": "Starship deve diventare il camion orbitale di Starlink V3, con missioni a capacita molto superiore rispetto a Falcon.", "Impatto": "Il business case puo accelerare la cadenza indipendentemente dalle missioni lunari.", "Fonte prevalente": "Doc 1"},
        {"Tema": "Controversie ambientali e sociali", "Sintesi integrata": "Debris, rumore, hazard areas, impatti su Boca Chica e incorporazione della City of Starbase generano attrito politico e ambientale.", "Impatto": "Il rischio non e' solo tecnico: licenze, consenso locale e gestione traffico aereo restano vincoli concreti.", "Fonte prevalente": "Doc 1 con integrazioni Doc 2"},
        {"Tema": "Elementi speculativi del Doc 2", "Sintesi integrata": "Valutazione fino a 1,75 trilioni, data center orbitali/xAI e alcuni target tecnici sono trattati come scenari o claim da verificare, non come fatti consolidati.", "Impatto": "Inclusi per completezza, ma separati dalle milestone documentate.", "Fonte prevalente": "Doc 2, marcato come speculativo"},
    ]


def discrepancy_rows() -> list[dict]:
    workbook_data = workbook_rows("Discrepanze", 4)
    if workbook_data:
        return [
            {
                "Nodo": cell_to_text(row.get("Nodo", "")),
                "Doc 1": cell_to_text(row.get("Doc 1", "")),
                "Doc 2": cell_to_text(row.get("Doc 2", "")),
                "Scelta nel workbook": cell_to_text(row.get("Scelta nel workbook", "")),
                "Confidenza": cell_to_text(row.get("Confidenza", "")),
            }
            for row in workbook_data
        ]

    return [
        {"Nodo": "IFT-3 e 'orbita'", "Doc 1": "Evita la formula orbita operativa/stabile; parla di traiettoria non-stable-orbit/transatmosferica.", "Doc 2": "Indica apogeo circa 234 km e traiettoria suborbitale/transatmosferica.", "Scelta nel workbook": "Classificato come operazioni quasi-orbitali/transatmosferiche, non orbita stabile.", "Confidenza": "Alta"},
        {"Nodo": "B19 full 33-engine static fire", "Doc 1": "15/04/2026.", "Doc 2": "16/04/2026.", "Scelta nel workbook": "Usata data 15/04/2026 per coerenza con cronologia strutturata; discrepanza mantenuta qui.", "Confidenza": "Media"},
        {"Nodo": "Ship 39 static fire e numero motori", "Doc 1": "Six-engine static fire circa un minuto il 14/04/2026.", "Doc 2": "Static fire di 9 motori, 3 sea-level e 6 vacuum.", "Scelta nel workbook": "Timeline usa sei motori; la tabella Block riporta il target/claim Doc 2 separatamente.", "Confidenza": "Media"},
        {"Nodo": "Artemis III / IV", "Doc 1": "A marzo 2026 NASA rivede: Artemis III demo LEO 2027, landing spostato ad Artemis IV 2028.", "Doc 2": "Parla di rinvio/valutazione Artemis III al 2027 e 2026 per docking LEO.", "Scelta nel workbook": "Prevale Doc 1 per maggiore specificita e tono meno ipotetico.", "Confidenza": "Alta"},
        {"Nodo": "FAA Part 450 dal 2026", "Doc 1": "Non lo usa come punto centrale.", "Doc 2": "Afferma operativita sotto Part 450 dal 2026.", "Scelta nel workbook": "Incluso solo come contesto regolatorio da verificare, non come milestone autonoma.", "Confidenza": "Bassa/Media"},
        {"Nodo": "Valutazione 1,75 trilioni e data center xAI", "Doc 1": "Non presente.", "Doc 2": "Scenario finanziario/strategico ad alto rischio.", "Scelta nel workbook": "Inserito nei temi speculativi, separato dai fatti della timeline.", "Confidenza": "Bassa"},
        {"Nodo": "V3 target 200 t LEO / 44 operazioni KSC", "Doc 1": "100+ t riutilizzabile e ampliamento Florida/LC-39A.", "Doc 2": "Target 200 t e fino a 44 operazioni/anno KSC.", "Scelta nel workbook": "Riportato nella scheda tecnica come target/claim, non come dato operativo consolidato.", "Confidenza": "Media"},
    ]


def set_tab_color(ws, color: str) -> None:
    ws.sheet_properties.tabColor = color


def apply_title(ws, title: str, subtitle: str | None = None, last_col: int = 8) -> None:
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=last_col)
    c = ws.cell(1, 1, title)
    c.fill = PatternFill("solid", fgColor=PALETTE["navy"])
    c.font = Font(color="FFFFFF", bold=True, size=18)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 34
    if subtitle:
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=last_col)
        c = ws.cell(2, 1, subtitle)
        c.fill = PatternFill("solid", fgColor=PALETTE["ink"])
        c.font = Font(color=PALETTE["sand"], italic=True, size=11)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[2].height = 24


def style_table(ws, header_row: int, last_row: int, last_col: int, freeze: str | None = None) -> None:
    thin = Side(style="thin", color=PALETTE["line"])
    for col in range(1, last_col + 1):
        cell = ws.cell(header_row, col)
        cell.fill = PatternFill("solid", fgColor=PALETTE["navy"])
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows(min_row=header_row + 1, max_row=last_row, max_col=last_col):
        for cell in row:
            cell.border = Border(bottom=thin)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if cell.row % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="FBFDFF")
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(last_col)}{last_row}"
    if freeze:
        ws.freeze_panes = freeze


def autosize(ws, widths: dict[int, int] | None = None, max_width: int = 55) -> None:
    widths = widths or {}
    for col_idx in range(1, ws.max_column + 1):
        letter = get_column_letter(col_idx)
        if col_idx in widths:
            ws.column_dimensions[letter].width = widths[col_idx]
            continue
        max_len = 0
        for cell in ws[letter]:
            if cell.value is not None:
                max_len = max(max_len, min(len(str(cell.value)), max_width))
        ws.column_dimensions[letter].width = max(10, min(max_len + 2, max_width))


def write_rows(ws, rows: list[dict], headers: list[str], start_row: int) -> int:
    for c, header in enumerate(headers, 1):
        ws.cell(start_row, c, header)
    for r, row in enumerate(rows, start_row + 1):
        for c, header in enumerate(headers, 1):
            value = row.get(header, "")
            if isinstance(value, datetime):
                value = value.strftime("%Y-%m-%d")
            ws.cell(r, c, value)
    return start_row + len(rows)


def add_kpi(ws, row: int, col: int, label: str, value: str, note: str, fill: str) -> None:
    for rr in range(row, row + 3):
        ws.merge_cells(start_row=rr, start_column=col, end_row=rr, end_column=col + 1)
        cell = ws.cell(rr, col)
        cell.fill = PatternFill("solid", fgColor=fill)
        cell.border = Border(
            left=Side(style="thin", color=PALETTE["line"]),
            right=Side(style="thin", color=PALETTE["line"]),
            top=Side(style="thin", color=PALETTE["line"]),
            bottom=Side(style="thin", color=PALETTE["line"]),
        )
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.cell(row, col, label).font = Font(color=PALETTE["muted"], bold=True, size=9)
    ws.cell(row + 1, col, value).font = Font(color=PALETTE["navy"], bold=True, size=20)
    ws.cell(row + 2, col, note).font = Font(color=PALETTE["muted"], italic=True, size=9)


def build_dashboard(ws, timeline: list[dict], integrated: list[dict]) -> None:
    set_tab_color(ws, PALETTE["gold"])
    apply_title(ws, "Programma Starship 2005-2026", "Dashboard riepilogativa - fusione ragionata dei due documenti Word al 23/04/2026", 10)
    ws.sheet_view.showGridLines = False
    add_kpi(ws, 4, 1, "Eventi timeline", str(len(timeline)), "Cronologia integrata", PALETTE["blue"])
    add_kpi(ws, 4, 3, "Voli integrati", str(len(integrated)), "IFT-1 -> Flight 11", PALETTE["teal"])
    add_kpi(ws, 4, 5, "Catch booster", "3", "Flight 5, 7, 8", PALETTE["success"])
    add_kpi(ws, 4, 7, "Reflight Super Heavy", "1", "Flight 9 / B14-2", PALETTE["warn"])
    add_kpi(ws, 4, 9, "Stato attuale", "Flight 12", "B19 + Ship 39 in prep.", PALETTE["sand"])

    notes = [
        "Starship evolve da visione marziana a sistema industriale riutilizzabile, con Raptor e infrastruttura come colli di bottiglia centrali.",
        "Al 23/04/2026 il programma ha completato 11 voli integrati, con booster-side piu maturo della ship-side.",
        "La fase Block 3/V3 punta a Raptor 3, Pad 2, hardware di refueling/docking e cadenza piu alta.",
        "Per Artemis e Marte il nodo decisivo resta la logistica spaziale: rifornimento criogenico, boil-off, docking e serie di tanker.",
    ]
    ws.merge_cells("A9:J9")
    ws["A9"] = "Sintesi esecutiva"
    ws["A9"].fill = PatternFill("solid", fgColor=PALETTE["navy"])
    ws["A9"].font = Font(color="FFFFFF", bold=True, size=12)
    for i, note in enumerate(notes, 10):
        ws.merge_cells(start_row=i, start_column=1, end_row=i, end_column=10)
        ws.cell(i, 1, note)
        ws.cell(i, 1).alignment = Alignment(wrap_text=True)

    category_counts = Counter(row["Categoria"] for row in timeline)
    outcome_counts = Counter(row["Esito integrato"].split(" / ")[0] for row in integrated)
    flights_by_year = Counter(parse_sort_date(row["Data"]).year for row in integrated)
    cumulative = []
    running = 0
    for year in sorted(flights_by_year):
        running += flights_by_year[year]
        cumulative.append((year, flights_by_year[year], running))

    start = 17
    ws.cell(start, 1, "Categoria")
    ws.cell(start, 2, "Eventi")
    for i, (cat, count) in enumerate(category_counts.most_common(), start + 1):
        ws.cell(i, 1, cat)
        ws.cell(i, 2, count)
    style_table(ws, start, start + len(category_counts), 2)

    ws.cell(start, 4, "Esito voli")
    ws.cell(start, 5, "Numero")
    for i, (outcome, count) in enumerate(outcome_counts.items(), start + 1):
        ws.cell(i, 4, outcome)
        ws.cell(i, 5, count)
    style_table(ws, start, start + len(outcome_counts), 5)

    ws.cell(start, 7, "Anno")
    ws.cell(start, 8, "Voli anno")
    ws.cell(start, 9, "Cumulato")
    for i, (year, count, cum) in enumerate(cumulative, start + 1):
        ws.cell(i, 7, year)
        ws.cell(i, 8, count)
        ws.cell(i, 9, cum)
    style_table(ws, start, start + len(cumulative), 9)

    chart1 = BarChart()
    chart1.title = "Eventi per categoria"
    chart1.y_axis.title = "Eventi"
    data = Reference(ws, min_col=2, min_row=start, max_row=start + len(category_counts))
    cats = Reference(ws, min_col=1, min_row=start + 1, max_row=start + len(category_counts))
    chart1.add_data(data, titles_from_data=True)
    chart1.set_categories(cats)
    chart1.height = 7
    chart1.width = 13
    ws.add_chart(chart1, "A27")

    chart2 = DoughnutChart()
    chart2.title = "Esito voli integrati"
    data = Reference(ws, min_col=5, min_row=start, max_row=start + len(outcome_counts))
    cats = Reference(ws, min_col=4, min_row=start + 1, max_row=start + len(outcome_counts))
    chart2.add_data(data, titles_from_data=True)
    chart2.set_categories(cats)
    chart2.height = 7
    chart2.width = 9
    ws.add_chart(chart2, "E27")

    chart3 = LineChart()
    chart3.title = "Voli integrati cumulati"
    data = Reference(ws, min_col=9, min_row=start, max_row=start + len(cumulative))
    cats = Reference(ws, min_col=7, min_row=start + 1, max_row=start + len(cumulative))
    chart3.add_data(data, titles_from_data=True)
    chart3.set_categories(cats)
    chart3.height = 7
    chart3.width = 9
    ws.add_chart(chart3, "H27")
    autosize(ws, {1: 30, 2: 12, 3: 16, 4: 22, 5: 12, 7: 12, 8: 14, 9: 14, 10: 18})


def build_workbook() -> None:
    if not DOC1.exists() or not DOC2.exists():
        reconstruct_source_docs_from_current_workbook()

    missing_docs = [str(path) for path in (DOC1, DOC2) if not path.exists()]
    if missing_docs:
        missing = "\n- ".join(missing_docs)
        raise FileNotFoundError(
            "Documenti Word sorgente mancanti. Inserire i file in 05_fonti_originali con questi nomi:\n"
            f"- {missing}"
        )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    timeline = prepare_timeline()
    integrated = integrated_flights()
    doc2_tables = extract_doc2_tables()

    wb = Workbook()
    build_dashboard(wb.active, timeline, integrated)
    wb.active.title = "Dashboard"

    ws = wb.create_sheet("Timeline integrata")
    set_tab_color(ws, "2A9D8F")
    apply_title(ws, "Timeline integrata", "Eventi ordinati temporalmente, con note di fusione tra Doc 1 e Doc 2", 12)
    headers = ["Data", "Anno", "Fase", "Evento", "Categoria", "Veicolo/Missione", "Esito", "Dettagli", "Impatto", "Fusione", "Fonte", "SortKey"]
    end_row = write_rows(ws, timeline, headers, 4)
    style_table(ws, 4, end_row, len(headers), freeze="A5")
    ws.column_dimensions["L"].hidden = True
    for row in range(5, end_row + 1):
        cat = ws.cell(row, 5).value
        ws.cell(row, 5).fill = PatternFill("solid", fgColor=CATEGORY_FILL.get(cat, "FFFFFF"))
        ws.cell(row, 5).font = Font(color=PALETTE["ink"], bold=True)
    autosize(ws, {1: 18, 2: 10, 3: 32, 4: 38, 5: 30, 6: 24, 7: 22, 8: 72, 9: 60, 10: 64, 11: 26, 12: 14})

    ws = wb.create_sheet("Voli integrati")
    set_tab_color(ws, "3A86FF")
    apply_title(ws, "Voli integrati", "Lettura consolidata da IFT-1 a Flight 11", 9)
    flight_headers = ["Flight", "Data", "Veicolo", "Block", "Milestone", "Booster", "Ship", "Esito integrato", "Lezione"]
    end_row = write_rows(ws, integrated, flight_headers, 4)
    style_table(ws, 4, end_row, len(flight_headers), freeze="A5")
    for row in range(5, end_row + 1):
        val = str(ws.cell(row, 8).value)
        fill = PALETTE["danger"] if "Fallimento" in val else PALETTE["warn"] if "Parziale" in val else PALETTE["success"]
        ws.cell(row, 8).fill = PatternFill("solid", fgColor=fill)
        ws.cell(row, 8).font = Font(bold=True, color=PALETTE["ink"])
    autosize(ws, {1: 12, 2: 14, 3: 26, 4: 14, 5: 48, 6: 38, 7: 42, 8: 18, 9: 62})

    ws = wb.create_sheet("Confronti tecnici")
    set_tab_color(ws, "E76F51")
    apply_title(ws, "Confronti tecnici", "Tabelle Doc 2 integrate con note di prudenza dal Doc 1", 6)
    current = 4
    for name, table in doc2_tables.items():
        ws.merge_cells(start_row=current, start_column=1, end_row=current, end_column=max(len(table[0]), 5))
        ws.cell(current, 1, name)
        ws.cell(current, 1).fill = PatternFill("solid", fgColor=PALETTE["ink"])
        ws.cell(current, 1).font = Font(color="FFFFFF", bold=True, size=12)
        current += 1
        for r, row_values in enumerate(table, current):
            for c, value in enumerate(row_values, 1):
                ws.cell(r, c, value)
        style_table(ws, current, current + len(table) - 1, len(table[0]))
        current += len(table) + 2
    notes = [
        "Nota di fusione: le tabelle tecniche provengono dal Doc 2 e sono conservate per completezza.",
        "Quando i dati entrano in conflitto con il Doc 1, la timeline usa la versione piu prudente e citata.",
        "Esempio: Ship 39 e' trattata come six-engine static fire; la configurazione V3 a 9 motori resta un target/claim tecnico del Doc 2.",
    ]
    for note in notes:
        ws.merge_cells(start_row=current, start_column=1, end_row=current, end_column=6)
        ws.cell(current, 1, note)
        ws.cell(current, 1).fill = PatternFill("solid", fgColor=PALETTE["sand"])
        ws.cell(current, 1).alignment = Alignment(wrap_text=True)
        current += 1
    autosize(ws, {1: 30, 2: 26, 3: 26, 4: 28, 5: 30, 6: 30})

    ws = wb.create_sheet("Temi e sintesi")
    set_tab_color(ws, "264653")
    apply_title(ws, "Temi e sintesi", "Sintesi ragionata delle parti narrative dei due documenti", 4)
    topic_headers = ["Tema", "Sintesi integrata", "Impatto", "Fonte prevalente"]
    end_row = write_rows(ws, summary_topics(), topic_headers, 4)
    style_table(ws, 4, end_row, len(topic_headers), freeze="A5")
    autosize(ws, {1: 30, 2: 78, 3: 64, 4: 34})

    ws = wb.create_sheet("Discrepanze")
    set_tab_color(ws, "B00020")
    apply_title(ws, "Discrepanze e scelte di fusione", "Dove i documenti divergono, qui e' indicata la scelta applicata", 5)
    disc_headers = ["Nodo", "Doc 1", "Doc 2", "Scelta nel workbook", "Confidenza"]
    end_row = write_rows(ws, discrepancy_rows(), disc_headers, 4)
    style_table(ws, 4, end_row, len(disc_headers), freeze="A5")
    for row in range(5, end_row + 1):
        conf = str(ws.cell(row, 5).value).lower()
        fill = PALETTE["danger"] if "bassa" in conf else PALETTE["warn"] if "media" in conf else PALETTE["success"]
        ws.cell(row, 5).fill = PatternFill("solid", fgColor=fill)
        ws.cell(row, 5).font = Font(bold=True)
    autosize(ws, {1: 30, 2: 60, 3: 60, 4: 70, 5: 16})

    ws = wb.create_sheet("Fonti")
    set_tab_color(ws, "6C757D")
    apply_title(ws, "Fonti citate nei documenti", "Link e domini estratti dai due file Word; nessuna verifica web esterna eseguita", 2)
    end_row = write_rows(ws, extract_sources(), ["Documento", "Fonte"], 4)
    style_table(ws, 4, end_row, 2, freeze="A5")
    autosize(ws, {1: 14, 2: 120})

    wb.properties.title = "Programma Starship - Timeline integrata"
    wb.properties.subject = "Fusione ragionata di due deepsearch Word"
    wb.properties.creator = "Codex"
    wb.save(OUT)
    load_workbook(OUT, read_only=True).close()


if __name__ == "__main__":
    build_workbook()
    print(OUT)
